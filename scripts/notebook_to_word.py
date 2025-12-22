#!/usr/bin/env python3
"""
Convert Jupyter Notebook to Word Document
"""

import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys


def convert_notebook_to_word(notebook_path: str, output_path: str = None):
    """
    Convert a Jupyter notebook to a Word document.
    
    Args:
        notebook_path: Path to the .ipynb file
        output_path: Path for the output .docx file (optional)
    """
    # Load notebook
    notebook_path = Path(notebook_path)
    if not notebook_path.exists():
        raise FileNotFoundError(f"Notebook not found: {notebook_path}")
    
    with open(notebook_path, 'r', encoding='utf-8') as f:
        notebook = json.load(f)
    
    # Create Word document
    doc = Document()
    
    # Set document title
    title = doc.add_heading(notebook_path.stem.replace('_', ' '), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    metadata = notebook.get('metadata', {})
    if metadata:
        doc.add_paragraph(f"Notebook: {notebook_path.name}")
        doc.add_paragraph(f"Cells: {len(notebook.get('cells', []))}")
        doc.add_paragraph("")
    
    # Process cells
    for idx, cell in enumerate(notebook.get('cells', []), 1):
        cell_type = cell.get('cell_type', 'unknown')
        source = cell.get('source', [])
        
        # Convert source to string if it's a list
        if isinstance(source, list):
            content = ''.join(source)
        else:
            content = source
        
        if not content.strip():
            continue
        
        if cell_type == 'markdown':
            # Process markdown cell
            lines = content.split('\n')
            for line in lines:
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('#### '):
                    doc.add_heading(line[5:], level=4)
                elif line.strip():
                    p = doc.add_paragraph(line)
                    
        elif cell_type == 'code':
            # Add code cell header
            heading = doc.add_heading(f'Code Cell {idx}', level=3)
            heading_format = heading.runs[0].font
            heading_format.color.rgb = RGBColor(0, 102, 204)
            
            # Add code with monospace formatting
            code_para = doc.add_paragraph()
            code_run = code_para.add_run(content)
            code_run.font.name = 'Courier New'
            code_run.font.size = Pt(9)
            code_para.paragraph_format.left_indent = Inches(0.5)
            code_para.paragraph_format.space_before = Pt(6)
            code_para.paragraph_format.space_after = Pt(6)
            
            # Add cell outputs if present
            outputs = cell.get('outputs', [])
            if outputs:
                doc.add_paragraph('Output:', style='Heading 4')
                for output in outputs:
                    output_type = output.get('output_type', '')
                    
                    if output_type == 'stream':
                        text = ''.join(output.get('text', []))
                        if text.strip():
                            output_para = doc.add_paragraph()
                            output_run = output_para.add_run(text)
                            output_run.font.name = 'Courier New'
                            output_run.font.size = Pt(8)
                            output_run.font.color.rgb = RGBColor(0, 128, 0)
                            output_para.paragraph_format.left_indent = Inches(0.75)
                    
                    elif output_type in ('execute_result', 'display_data'):
                        data = output.get('data', {})
                        if 'text/plain' in data:
                            text = ''.join(data['text/plain'])
                            if text.strip():
                                result_para = doc.add_paragraph()
                                result_run = result_para.add_run(text)
                                result_run.font.name = 'Courier New'
                                result_run.font.size = Pt(8)
                                result_run.font.color.rgb = RGBColor(128, 0, 128)
                                result_para.paragraph_format.left_indent = Inches(0.75)
                    
                    elif output_type == 'error':
                        error_name = output.get('ename', 'Error')
                        error_value = output.get('evalue', '')
                        error_para = doc.add_paragraph()
                        error_run = error_para.add_run(f"{error_name}: {error_value}")
                        error_run.font.name = 'Courier New'
                        error_run.font.size = Pt(8)
                        error_run.font.color.rgb = RGBColor(255, 0, 0)
                        error_para.paragraph_format.left_indent = Inches(0.75)
            
            # Add spacing after code cell
            doc.add_paragraph()
    
    # Save document
    if output_path is None:
        output_path = notebook_path.with_suffix('.docx')
    else:
        output_path = Path(output_path)
    
    doc.save(str(output_path))
    print(f"✓ Converted: {notebook_path.name}")
    print(f"✓ Saved to: {output_path}")
    print(f"✓ Total cells processed: {len(notebook.get('cells', []))}")
    
    return output_path


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python notebook_to_word.py <notebook.ipynb> [output.docx]")
        sys.exit(1)
    
    notebook_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None
    
    try:
        convert_notebook_to_word(notebook_path, output_path)
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)
